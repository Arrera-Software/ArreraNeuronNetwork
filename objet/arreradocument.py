from docx import Document
from odf import text, opendocument ,teletype
from odf.opendocument import load

class CArreraDocx :
    def __init__(self,fileName:str) :
        if ("docx" in fileName) :
            self.__document = Document()
            self.__fileName = fileName
            self.__typeFile = "docx"
            self.__init = True
        else :
            if ("odt"in fileName):
                self.__document = opendocument.OpenDocumentText()
                self.__fileName = fileName
                self.__typeFile = "odt"
                self.__init = True
            else :
                self.__fileName = fileName
                self.__init = False
    
    def write(self,content:str):
        if (self.__init==True):
            if (self.__typeFile=="docx"):
                for line in content.split('\n'):
                    self.__document.add_paragraph(line)
                self.__document.save(self.__fileName)
                return True
            else :
                if (self.__typeFile=="odt"):
                    for line in content.split('\n'):
                        p = text.P()
                        p.addText(line)
                        self.__document.text.addElement(p)
                    self.__document.save(self.__fileName)
                    return True
        else :
            return False
    
    def read(self):
        if (self.__init==True):
            if (self.__typeFile=="docx"):
                self.__document = Document(self.__fileName)
                texte = []
                for paragraph in self.__document.paragraphs:
                    texte.append(paragraph.text)
                return '\n'.join(texte)
            else :
                if (self.__typeFile=="odt"):
                    self.__document = load(self.__fileName)
                    content = ""
                    for para in self.__document.getElementsByType(text.P):
                        content += teletype.extractText(para)+"\n"
                    return content
        else :
            return "error"